from fpdf import FPDF
from docx import Document
from docx.enum.text import WD_BREAK
import os

def create_samples():
    os.makedirs("samples", exist_ok=True)
    
    # 1. Create Multi-page TXT files (just long content)
    long_text = "Detailed Project Documentation\n" + "="*30 + "\n\n"
    for i in range(1, 21):
        long_text += f"Section {i}: Deep Dive into Topic {i}\n"
        long_text += "This section provides an in-depth analysis of the subject matter at hand. We explore various angles and data points to ensure a comprehensive understanding. The importance of this specific topic cannot be understated in the context of the overall project goals. We have conducted extensive research and interviewed several key stakeholders to compile this information. It serves as a foundational element for the subsequent phases of our strategy. Continuous monitoring and evaluation will be required to maintain the relevance of these findings over the coming months. We anticipate some challenges in implementation but remain confident in our methodology.\n\n"
    
    with open("samples/long_project_doc.txt", "w") as f:
        f.write(long_text)
    
    with open("samples/extensive_notes.txt", "w") as f:
        f.write(long_text.replace("Topic", "Note"))

    # 2. Create Multi-page DOCX files
    doc1 = Document()
    doc1.add_heading('Technical Specification v5.1', 0)
    for i in range(1, 6):
        doc1.add_heading(f'Chapter {i}: Architectural Components', level=1)
        doc1.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.')
        doc1.add_paragraph('Detailed system diagrams and flowcharts are essential for the engineering team to visualize the data flow. We utilize a microservices-based approach with containerized deployments to ensure scalability and fault tolerance. Each service communicates via asynchronous message queues to decouple the components and improve reliability.')
        doc1.add_page_break()
    doc1.save("samples/multi_page_tech.docx")

    doc2 = Document()
    doc2.add_heading('Operational Manual 2026', 0)
    for i in range(1, 6):
        doc2.add_heading(f'Standard Operating Procedure {i}', level=1)
        doc2.add_paragraph('All staff members must adhere to these procedures to ensure high quality and safety standards. Regular audits will be conducted to verify compliance. Training sessions are provided quarterly to keep everyone updated on the latest best practices and regulatory changes.')
        doc2.add_paragraph('In case of any discrepancies or emergencies, follow the escalation matrix provided in the appendix. Prompt reporting is mandatory for all incidents. We maintain a non-punitive reporting culture to encourage transparency and collective improvement across all departments.')
        doc2.add_page_break()
    doc2.save("samples/multi_page_ops.docx")

    # 3. Create Multi-page PDF files
    class PDF(FPDF):
        def header(self):
            self.set_font('helvetica', 'B', 12)
            self.cell(0, 10, 'White Paper: Future Technologies', 0, 1, 'C')
            self.ln(5)

    pdf1 = PDF()
    pdf1.set_auto_page_break(auto=True, margin=15)
    for i in range(1, 6):
        pdf1.add_page()
        pdf1.set_font('helvetica', 'B', 14)
        pdf1.cell(0, 10, f"Page {i}: Deep Analysis", ln=True)
        pdf1.set_font('helvetica', '', 11)
        long_pdf_text = """This document serves as a comprehensive guide to the emerging technological trends that are set to redefine the global landscape by 2030. We examine the convergence of biotechnology, nanotechnology, and quantum computing. 

The first major trend is the decentralization of manufacturing through advanced 3D printing. This will lead to localized production hubs, reducing the carbon footprint of global shipping. 

Secondly, the rise of sovereign AI models will allow nations to maintain data sovereignty while benefiting from cutting-edge machine learning capabilities. 

Lastly, the fusion of augmented reality and real-time data visualization will transform fields like medicine, construction, and education, providing experts with hands-on digital assistance in their physical environments.
""" * 4 # Repeat to fill some space
        pdf1.multi_cell(0, 10, long_pdf_text)
    pdf1.output("samples/multi_page_future_tech.pdf")

    pdf2 = PDF()
    for i in range(1, 6):
        pdf2.add_page()
        pdf2.set_font('helvetica', 'B', 14)
        pdf2.cell(0, 10, f"Section {i}: Market Forecast", ln=True)
        pdf2.set_font('helvetica', '', 11)
        pdf_text2 = """The following analysis provides a deep dive into the primary drivers of growth in the sustainable energy sector. We have synthesized data from top environmental think tanks and financial institutions.

Key findings include a project 200% increase in utility-scale battery deployments by the end of the decade. This is driven by falling lithium-ion battery costs and favorable government subsidies in major economies.

We also anticipate a shift in consumer behavior toward sustainable living, which will create new markets for green home technologies and electric vehicle infrastructure. 

The financial sector will play a crucial role by providing the necessary capital through green bonds and impact investment funds. This systemic change is essential for meeting global climate targets.
""" * 4
        pdf2.multi_cell(0, 10, pdf_text2)
    pdf2.output("samples/multi_page_energy_forecast.pdf")

    print("Large multi-page sample files created in /samples directory.")

if __name__ == "__main__":
    create_samples()
